"""DOCX extraction using python-docx."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from ingest.extractors.base import ExtractedDocument


class DOCXExtractor:
    """Extract accepted text and heading sections from Word documents."""

    def extract(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        """Extract paragraphs, core metadata, and heading-delimited sections."""

        document = Document(BytesIO(file_bytes))
        sections: list[dict[str, str]] = []
        heading = "Document"
        buffer: list[str] = []
        text_parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.startswith("Heading"):
                if buffer:
                    sections.append({"heading": heading, "content": "\n".join(buffer).strip()})
                heading = text
                buffer = []
            else:
                buffer.append(text)
            text_parts.append(text)

        if buffer:
            sections.append({"heading": heading, "content": "\n".join(buffer).strip()})

        raw_text = "\n\n".join(text_parts)
        props = document.core_properties
        title = props.title or filename
        return ExtractedDocument(
            title=title,
            raw_text=raw_text,
            metadata={
                "author": props.author or None,
                "date_created": props.created.isoformat() if props.created else None,
                "date_modified": props.modified.isoformat() if props.modified else None,
                "page_count": None,
                "slide_count": None,
                "word_count": len(raw_text.split()),
            },
            sections=sections or [{"heading": "Document", "content": raw_text}],
        )
