"""End-to-end ingestion pipeline tests."""

from __future__ import annotations

from docx import Document

from config import Settings
from db.database import Database
from ingest.pipeline import IngestPipeline


def test_pipeline_ingests_local_docx(tmp_path):
    docs = tmp_path / "docs"
    docs.mkdir()
    doc = Document()
    doc.add_heading("Avionics Project", level=1)
    doc.add_paragraph("Embedded avionics research with FPGA prototype.")
    doc.save(docs / "project.docx")

    settings = Settings(local_docs_path=docs, db_path=tmp_path / "kb.db", index_path=tmp_path / "bm25.pkl")
    with Database(settings.db_path) as database:
        database.init_schema()
        summary = IngestPipeline(settings, database).run("local")
        stats = database.stats()

    assert summary.processed == 1
    assert stats["projects"] == 1
    assert stats["chunks"] >= 1


def test_pipeline_skips_duplicate_local_docx_on_second_run(tmp_path):
    docs = tmp_path / "docs"
    docs.mkdir()
    doc = Document()
    doc.add_heading("Radar Project", level=1)
    doc.add_paragraph("FPGA radar processing prototype.")
    doc.save(docs / "radar.docx")

    settings = Settings(local_docs_path=docs, db_path=tmp_path / "kb.db", index_path=tmp_path / "bm25.pkl")
    with Database(settings.db_path) as database:
        database.init_schema()
        pipeline = IngestPipeline(settings, database)
        first = pipeline.run("local")
        second = pipeline.run("local")
        stats = database.stats()

    assert first.processed == 1
    assert second.processed == 0
    assert second.skipped == 1
    assert second.failed == 0
    assert stats["projects"] == 1
